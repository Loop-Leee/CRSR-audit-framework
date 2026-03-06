"""DOCX 批注写入工具（基于 python-docx）。"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from .logger import Logger

try:
    from docx import Document  # type: ignore[import-not-found]
    from docx.table import Table  # type: ignore[import-not-found]
    from docx.text.paragraph import Paragraph  # type: ignore[import-not-found]
except Exception:  # noqa: BLE001
    Document = None  # type: ignore[assignment]
    Table = None  # type: ignore[assignment]
    Paragraph = None  # type: ignore[assignment]


@dataclass(frozen=True, slots=True)
class DocxComment:
    """单条批注。"""

    risk_id: str
    anchor_offset: int
    span: str
    content: str


@dataclass(frozen=True, slots=True)
class DocxCommentWriteResult:
    """批注写入结果。"""

    written_count: int
    skipped_count: int
    written_mask: list[bool]


@dataclass(frozen=True, slots=True)
class _ParagraphMeta:
    index: int
    start: int
    end: int
    text: str
    paragraph: Any


@dataclass(frozen=True, slots=True)
class _SpanAnchor:
    paragraph: Any
    start_char: int
    end_char: int


@dataclass(frozen=True, slots=True)
class _RunSegment:
    run: Any
    start: int
    end: int


def ensure_docx_comment_support() -> None:
    """校验当前环境是否具备 DOCX 批注写入能力。"""

    if Document is None:
        raise RuntimeError("未安装 python-docx。请在环境中安装 `python-docx>=1.2.0`。")
    probe = Document()
    if not hasattr(probe, "add_comment"):
        raise RuntimeError("当前 python-docx 版本不支持 comments API。请升级到 `python-docx>=1.2.0`。")


def add_comments_to_docx(
    *,
    docx_path: Path,
    comments: list[DocxComment],
    logger: Logger | None = None,
) -> DocxCommentWriteResult:
    """向 `.docx` 文档追加批注。

    Args:
        docx_path: 目标 `.docx` 文件路径。
        comments: 待写入批注列表。
        logger: 可选日志对象。

    Returns:
        DocxCommentWriteResult: 写入统计与逐条状态。

    Raises:
        FileNotFoundError: 文件不存在时抛出。
        ValueError: 输入文件不是 `.docx` 时抛出。
        RuntimeError: 缺少 python-docx 或版本不支持 comments API 时抛出。
    """

    if not docx_path.exists():
        raise FileNotFoundError(f"docx 文件不存在: {docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"仅支持 .docx 文件批注: {docx_path}")
    if not comments:
        return DocxCommentWriteResult(written_count=0, skipped_count=0, written_mask=[])
    ensure_docx_comment_support()

    document = Document(str(docx_path))

    paragraph_meta = _build_paragraph_meta(document=document)

    written_mask: list[bool] = [False] * len(comments)
    for index, item in enumerate(comments):
        anchor = _select_span_anchor(
            paragraph_meta=paragraph_meta,
            anchor=item.anchor_offset,
            span=item.span,
        )
        if anchor is None:
            continue
        if not _add_comment_to_span(
            document=document,
            paragraph=anchor.paragraph,
            start_char=anchor.start_char,
            end_char=anchor.end_char,
            text=item.content,
        ):
            continue
        written_mask[index] = True

    written_count = sum(1 for flag in written_mask if flag)
    skipped_count = len(comments) - written_count
    if written_count <= 0:
        if logger:
            logger.info(f"docx 批注跳过：无可写入锚点。file={docx_path}, comment_count={len(comments)}")
        return DocxCommentWriteResult(
            written_count=written_count,
            skipped_count=skipped_count,
            written_mask=written_mask,
        )

    document.save(str(docx_path))
    if logger:
        logger.info(
            "docx 批注完成: file=%s, written=%s, skipped=%s"
            % (docx_path, written_count, skipped_count)
        )
    return DocxCommentWriteResult(
        written_count=written_count,
        skipped_count=skipped_count,
        written_mask=written_mask,
    )


def _build_paragraph_meta(*, document: Any) -> list[_ParagraphMeta]:
    paragraphs = _iter_paragraphs_in_order(document)
    cursor = 0
    result: list[_ParagraphMeta] = []
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text or ""
        start = cursor
        end = start + len(text)
        result.append(
            _ParagraphMeta(
                index=index,
                start=start,
                end=end,
                text=text,
                paragraph=paragraph,
            )
        )
        cursor = end + 1
    return result


def _iter_paragraphs_in_order(document: Any) -> list[Any]:
    if Paragraph is None or Table is None:
        return list(getattr(document, "paragraphs", []))

    result: list[Any] = []
    visited_cells: set[int] = set()
    _append_paragraphs_from_parent(
        parent=document,
        parent_element=document.element.body,
        result=result,
        visited_cells=visited_cells,
    )
    return result


def _append_paragraphs_from_parent(
    *,
    parent: Any,
    parent_element: Any,
    result: list[Any],
    visited_cells: set[int],
) -> None:
    for child in parent_element.iterchildren():
        tag_name = child.tag.rsplit("}", 1)[-1]
        if tag_name == "p":
            result.append(Paragraph(child, parent))
            continue
        if tag_name != "tbl":
            continue
        table = Table(child, parent)
        for row in table.rows:
            for cell in row.cells:
                cell_key = id(cell._tc)
                if cell_key in visited_cells:
                    continue
                visited_cells.add(cell_key)
                _append_paragraphs_from_parent(
                    parent=cell,
                    parent_element=cell._tc,
                    result=result,
                    visited_cells=visited_cells,
                )


def _select_span_anchor(
    paragraph_meta: list[_ParagraphMeta],
    *,
    anchor: int,
    span: str,
) -> _SpanAnchor | None:
    if not paragraph_meta:
        return None

    if span:
        candidates: list[tuple[int, int, _ParagraphMeta, int]] = []
        for item in paragraph_meta:
            search_from = 0
            while True:
                local_index = item.text.find(span, search_from)
                if local_index < 0:
                    break
                global_start = item.start + local_index
                distance = abs(global_start - anchor)
                candidates.append((distance, global_start, item, local_index))
                search_from = local_index + 1
        if candidates:
            candidates.sort(key=lambda row: (row[0], row[1], row[2].index))
            _, _, meta, local_start = candidates[0]
            return _SpanAnchor(
                paragraph=meta.paragraph,
                start_char=local_start,
                end_char=local_start + len(span),
            )

    selected = _select_paragraph_meta_by_anchor(paragraph_meta=paragraph_meta, anchor=anchor, span=span)
    if selected is None:
        return None
    if span:
        local_start = selected.text.find(span)
        if local_start >= 0:
            return _SpanAnchor(
                paragraph=selected.paragraph,
                start_char=local_start,
                end_char=local_start + len(span),
            )
    if selected.text:
        relative = max(0, min(len(selected.text) - 1, anchor - selected.start))
        return _SpanAnchor(
            paragraph=selected.paragraph,
            start_char=relative,
            end_char=relative + 1,
        )
    return _SpanAnchor(paragraph=selected.paragraph, start_char=0, end_char=0)


def _select_paragraph_meta_by_anchor(
    *,
    paragraph_meta: list[_ParagraphMeta],
    anchor: int,
    span: str,
) -> _ParagraphMeta | None:
    for item in paragraph_meta:
        if item.start <= anchor <= item.end:
            return item

    if span:
        for item in paragraph_meta:
            if span in item.text:
                return item

    nearest = min(paragraph_meta, key=lambda item: abs(item.start - anchor))
    return nearest


def _add_comment_to_span(
    *,
    document: Any,
    paragraph: Any,
    start_char: int,
    end_char: int,
    text: str,
) -> bool:
    segments = _build_run_segments(paragraph)
    if not segments:
        run = paragraph.add_run("\u200b")
        _add_comment_compat(document=document, runs=[run], text=text)
        return True

    span_start = max(0, start_char)
    span_end = max(span_start + 1, end_char)
    target_runs = [
        segment.run
        for segment in segments
        if segment.end > span_start and segment.start < span_end
    ]
    if not target_runs:
        nearest = min(segments, key=lambda segment: abs(segment.start - span_start))
        target_runs = [nearest.run]

    _add_comment_compat(document=document, runs=target_runs, text=text)
    return True


def _build_run_segments(paragraph: Any) -> list[_RunSegment]:
    cursor = 0
    segments: list[_RunSegment] = []
    for run in paragraph.runs:
        run_text = run.text or ""
        if not run_text:
            continue
        start = cursor
        end = start + len(run_text)
        segments.append(_RunSegment(run=run, start=start, end=end))
        cursor = end
    return segments


def _add_comment_compat(*, document: Any, runs: list[Any], text: str) -> None:
    add_comment = getattr(document, "add_comment")
    if not runs:
        raise RuntimeError("python-docx add_comment 调用失败: 缺少可用 run。")
    run_seq = list(runs)

    attempts = (
        lambda: add_comment(runs=run_seq, text=text, author="", initials=""),
        lambda: add_comment(run_seq, text=text, author="", initials=""),
        lambda: add_comment(run_seq, text, "", ""),
        lambda: add_comment(run_seq[0], text=text, author="", initials=""),
        lambda: add_comment(run_seq[0], text, "", ""),
    )

    last_error: Exception | None = None
    for attempt in attempts:
        try:
            attempt()
            return
        except TypeError as exc:
            last_error = exc
            continue

    raise RuntimeError(f"python-docx add_comment 调用失败: {last_error}")
